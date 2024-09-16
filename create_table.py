import os
import pathlib
import math
import threading
import pandas as pd
import docx
from docx.enum.section import WD_ORIENTATION
from docxtpl import DocxTemplate
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt
from docx.enum.table import WD_ROW_HEIGHT_RULE

import datetime
import traceback

from PyQt5.QtCore import QThread, pyqtSignal
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Font, Border, Side


class CreateTable(QThread):  # Если требуется вставить колонтитулы
    progress = pyqtSignal(int)  # Сигнал для прогресс бара
    status = pyqtSignal(str)  # Сигнал для статус бара
    messageChanged = pyqtSignal(str, str)
    errors = pyqtSignal()

    def __init__(self, incoming_data):  # Список переданных элементов.
        QThread.__init__(self)
        self.path_file = incoming_data['path_file']
        self.finish_path = incoming_data['finish_path']
        self.file_name = incoming_data['file_name']
        self.queue = incoming_data['queue']
        self.logging = incoming_data['logging']
        self.event = threading.Event()
        self.event.set()

    def set_vertical_cell_direction(self, cell: _Cell, direction: str):
        # direction: tbRl -- top to bottom, btLr -- bottom to top
        assert direction in ("tbRl", "btLr")
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        textDirection = OxmlElement('w:textDirection')
        textDirection.set(qn('w:val'), direction)  # btLr tbRl
        tcPr.append(textDirection)

    def run(self):
        try:
            progress = 0
            self.logging.info("\n*******************************************************************************\n")
            self.logging.info('Начинаем создание таблицы')
            self.status.emit('Считываем значения из файла')
            self.progress.emit(progress)
            time_start = datetime.datetime.now()
            index = ['numSet', 'numTs', 'nameTs', 'firm', 'model', 'sn', 'quant',
                     'secTs', 'secRom', 'first', 'second', 'secondK', 'mod', 'dat', 'snMni',
                     'textMni', 'plat', 'textPlat']
            name_col = ['№ комплекта', '№ ТС в комплекте', 'Наименование ТС', 'Фирма-производитель', 'Модель (Тип)',
                        'Заводской номер', 'Кол-во, ТС  шт.', 'Степень секретности информации, обрабатываемой ТС',
                        'Категория помещений, в котором установлено ТС', 'СЗЗ-1', 'СЗЗ-2', 'СЗЗ-2к',
                        'Наименование и модель удалённого модуля позволяющего организовать канал обмена информацией',
                        'Тип, наименование датчиков физических полей в составе ТС',
                        'Тип, модель, заводской  номер подлежащего регистрации несъёмного МНИ',
                        'Краткое описание места размещения несъёмного МНИ',
                        'Платы содержащие элементы накопления и хранения информации',
                        'Краткое описание места размещения плат, содержащих элементы накопления и хранения информации']
            name_1_col = ['Тип и кол-во СЗЗ, нанесенных на каждое ТС', 'Несъёмные МНИ в составе ТС',
                          'Элемент накопления и хранения информации в составе ТС']
            df = pd.read_csv(self.path_file, delimiter='|', encoding='ANSI', header=None)
            serial_number = ''
            incoming_errors = []
            if df[0].isnull().any():
                for ind, (first, second) in enumerate(zip(df[0].to_numpy(), df[1].to_numpy())):
                    if second == 1:
                        serial_number = first if pd.isna(first) is False else ''  # как вернуться к предыдущей?
                    if pd.isna(first):
                        incoming_errors.append(str(ind + 1))
                        df.loc[ind, 0] = serial_number
            # Здесь надо добавить предупреждение, чтобы при отсутсвии серийника в первой строке не падала программа
            if (int(math.log10(df.loc[0, 0]))+1) == 8:  # Для сверки номеров. Если это серийник - преобразование к int
                index_null = [str(ind + 1) for ind, val in enumerate(df[0].to_numpy().tolist()) if len(str(val)) == 0]
                if index_null:
                    self.logging.info("Есть строки с пустыми sn: " + ', '.join(index_null))
                    self.status.emit('Ошибка')  # Посылаем значние если готово
                    self.queue.put({'errors': index_null, 'title': 'Ошибки в загруженных данных,'
                                                                   ' номера строк с пропуском sn:\n'})
                    self.errors.emit()
                    self.progress.emit(0)
                    return
                df = df.astype({0: int})
                df = df.astype({0: str})
                df[0] = ['00' + element for element in df[0]]
            table_contents = []
            len_rows = {}
            number_set = []
            number_set_val = 1000000000
            progress += 15
            self.logging.info('Преобразовываем df и формируем таблицу высоты, если требуется')
            self.status.emit('Преобразование данных')
            self.progress.emit(progress)
            previous_val = ''
            max_size_col_excel = 10
            for ind_row, row in enumerate(df.itertuples()):
                len_string = 0
                if row[1] != number_set_val:
                    number_set_val = row[1]
                    number_set.append(ind_row + 2)
                list_val = [j for i, j in enumerate(row) if i > 0]
                dict_val = {}
                for ind, val in enumerate(list_val):
                    if ind == 1:
                        if pd.isna(val):
                            dict_val[index[ind]] = previous_val
                        else:
                            dict_val[index[ind]] = int(val) if val % 1 == 0 else val
                            previous_val = dict_val[index[ind]]
                    elif 8 < ind < 12 and (pd.isna(val) or val == '-'):
                        dict_val[index[ind]] = 0
                    elif ind >= 12 and (pd.isna(val) or val == '-'):
                        dict_val[index[ind]] = 'Отсутствуют'
                    else:
                        if ind == 6 or 8 < ind < 12:
                            dict_val[index[ind]] = int(val)
                        else:
                            dict_val[index[ind]] = val
                    if ind >= 12 and isinstance(val, str) and len(val) > len_string:
                        len_string = len(val)
                table_contents.append(dict_val)
                index_str = ind_row + 2
                if 60 < len_string <= 80:
                    len_rows[index_str] = 7
                    if max_size_col_excel <= 15:
                        max_size_col_excel = 15
                elif 80 < len_string <= 100:
                    len_rows[index_str] = 8
                    if max_size_col_excel <= 20:
                        max_size_col_excel = 20
                elif 100 < len_string <= 120:
                    len_rows[index_str] = 9
                    if max_size_col_excel <= 25:
                        max_size_col_excel = 25
                elif 120 < len_string <= 150:
                    len_rows[index_str] = 10
                    if max_size_col_excel <= 30:
                        max_size_col_excel = 30
                elif 150 < len_string <= 170:
                    len_rows[index_str] = 12
                    if max_size_col_excel <= 35:
                        max_size_col_excel = 35
                elif 170 < len_string <= 200:
                    len_rows[index_str] = 14
                    if max_size_col_excel <= 40:
                        max_size_col_excel = 40
                elif 200 < len_string <= 250:
                    len_rows[index_str] = 16
                    if max_size_col_excel <= 45:
                        max_size_col_excel = 45
                elif 250 < len_string <= 300:
                    len_rows[index_str] = 18
                    if max_size_col_excel <= 50:
                        max_size_col_excel = 50
                elif len_string > 300:
                    len_rows[index_str] = 25
                    if max_size_col_excel <= 55:
                        max_size_col_excel = 55
            number_set.append(len(df) + 2)
            self.logging.info('Удаляем отчёт с таким же именем, если он есть')
            while True:
                try:
                    os.remove(pathlib.Path(self.finish_path, f'{self.file_name}.docx'))
                    break
                except PermissionError:
                    self.messageChanged.emit('Вопрос?', f'Файл «{self.file_name}.docx» в проверямой папке должен быть'
                                                        f' перезаписан. При необходимости сохраните файл в другое место'
                                                        f' и закройте его. После этого нажмите «Да» для продолжения'
                                                        f' или «Нет» для прерывания')
                    self.event.clear()
                    self.event.wait()
                    if self.queue.get_nowait() is False:
                        self.logging.warning('Процесс прерван пользователем при создании docx файла')
                        self.status.emit('Прервано пользователем')
                        self.progress.emit(0)
                        return
                except FileNotFoundError:
                    break
            progress += 15
            self.status.emit('Создаем таблицу excel')
            self.logging.info('Создаем таблицу excel')
            # Новый отчёт excel
            while True:
                try:
                    if pathlib.Path(self.finish_path, f'{self.file_name}.xlsx').exists():
                        os.remove(pathlib.Path(self.finish_path, f'{self.file_name}.xlsx'))
                    with pd.ExcelWriter(pathlib.Path(self.finish_path, f'{self.file_name}.xlsx'),
                                        engine='openpyxl', mode='w') as writer:
                        df.to_excel(writer, sheet_name=self.file_name, index=False, header=False, startrow=2)
                    break
                except PermissionError:
                    self.messageChanged.emit('Вопрос?', f'Файл «{self.file_name}.xlsx» в проверямой папке должен быть'
                                                        f' перезаписан. При необходимости сохраните файл в другое место'
                                                        f' и закройте его. После этого нажмите «Да» для продолжения'
                                                        f' или «Нет» для прерывания')
                    self.event.clear()
                    self.event.wait()
                    if self.queue.get_nowait() is False:
                        self.logging.warning('Процесс прерван пользователем при создании excel файла')
                        self.status.emit('Прервано пользователем')
                        self.progress.emit(0)
                        return
            wb = load_workbook(pathlib.Path(self.finish_path, f'{self.file_name}.xlsx'))
            ws = wb.active
            thin = Side(border_style="thin", color="000000")
            col_dimension = [5, 5, 25, 15, 15, 18, 5, 5, 5, 4, 4, 4, 5, 5, 6, 6] + [max_size_col_excel]*2
            col_name = [chr(i) for i in range(65, 84)]
            for i in range(1, 19):
                ws.column_dimensions[col_name[i - 1]].width = col_dimension[i - 1]
                ws.cell(1, i).font = Font(name="Times New Roman", size="10")
                ws.cell(2, i).font = Font(name="Times New Roman", size="10")
                ws.cell(1, i).border = Border(top=thin, left=thin, right=thin, bottom=thin)
                ws.cell(2, i).border = Border(top=thin, left=thin, right=thin, bottom=thin)
                ws.cell(1, i).alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                ws.cell(2, i).alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                if i in [10, 11, 12, 15, 16]:
                    ws.cell(2, i).alignment = Alignment(horizontal="center", vertical="center", text_rotation=90,
                                                        wrap_text=True)
                if i not in [3, 4, 5, 6, 10, 11, 12, 15, 16, 17, 18]:
                    ws.cell(1, i).alignment = Alignment(horizontal="center", vertical="center", text_rotation=90,
                                                        wrap_text=True)
                if i < 10 or i in [13, 14]:
                    ws.merge_cells(start_row=1, start_column=i, end_row=2, end_column=i)
                    ws.cell(1, i).value = name_col[i - 1]
                elif i == 10:
                    ws.merge_cells(start_row=1, start_column=i, end_row=1, end_column=12)
                    ws.cell(2, i).value = name_col[i - 1]
                    ws.cell(1, i).value = name_1_col[0]
                elif i in [11, 12]:
                    ws.cell(2, i).value = name_col[i - 1]
                elif i == 15:
                    ws.merge_cells(start_row=1, start_column=i, end_row=1, end_column=16)
                    ws.cell(2, i).value = name_col[i - 1]
                    ws.cell(1, i).value = name_1_col[1]
                elif i == 17:
                    ws.merge_cells(start_row=1, start_column=i, end_row=1, end_column=18)
                    ws.cell(2, i).value = name_col[i - 1]
                    ws.cell(1, i).value = name_1_col[2]
                else:
                    ws.cell(2, i).value = name_col[i - 1]
            ws.row_dimensions[1].height = 50
            ws.row_dimensions[2].height = 250
            number_start = False
            number_finish = False
            for number in number_set:
                if number_start is False:
                    number_start = number + 1
                elif number_finish is False:
                    number_finish = number
                    ws.merge_cells(start_row=number_start, start_column=1, end_row=number_finish, end_column=1)
                    ws.cell(number_start, 1).font = Font(name="Times New Roman", size="10")
                    ws.cell(number_start, 1).border = Border(top=thin, left=thin, right=thin, bottom=thin)
                    ws.cell(number_start, 1).alignment = Alignment(horizontal="center", vertical="center",
                                                                   text_rotation=90, wrap_text=True)
                    second_number = number_start
                    for j in range(number_start, number_finish + 1):
                        ws.cell(j, 1).border = Border(top=thin, left=thin, right=thin, bottom=thin)
                        ws.cell(j, 2).border = Border(top=thin, left=thin, right=thin, bottom=thin)
                        ws.cell(j, 2).alignment = Alignment(horizontal="center", vertical="center")
                        ws.cell(j, 2).font = Font(name="Times New Roman", size="10")
                        if ws.cell(j, 2).value is not None:
                            if second_number != j and second_number != j - 1:
                                ws.merge_cells(start_row=second_number, start_column=2, end_row=j - 1, end_column=2)
                            second_number = j
                    number_start = number + 1
                    number_finish = False
            for i in range(3, df.shape[0] + 3):
                for j in range(3, 13):
                    ws.cell(i, j).font = Font(name="Times New Roman", size="10")
                    ws.cell(i, j).border = Border(top=thin, left=thin, right=thin, bottom=thin)
                    ws.cell(i, j).alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            row_dimension = {250: 70, 500: 140, 700: 200, 1000: 270, 10000: 400}
            for r, row in enumerate(ws[f'M3:R{df.shape[0] + 2}']):
                text_len = 0
                for cell in row:
                    if len(cell.value) > text_len:
                        text_len = len(cell.value)
                for i, cell in enumerate(row):
                    cell.font = Font(name="Times New Roman", size="10")
                    if i <= 3:
                        cell.alignment = Alignment(horizontal="center", vertical="center", text_rotation=90,
                                                   wrap_text=True)
                    else:
                        if text_len > 250:
                            cell.alignment = Alignment(horizontal="center", vertical="center", text_rotation=90,
                                                       wrap_text=True)
                        else:
                            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                    cell.border = Border(top=thin, left=thin, right=thin, bottom=thin)
                for size in row_dimension:
                    if text_len <= size:
                        ws.row_dimensions[r + 3].height = row_dimension[size]
                        break
            wb.save(pathlib.Path(self.finish_path, f'{self.file_name}.xlsx'))
            progress += 5
            self.progress.emit(progress)
            self.logging.info('Создаем шаблон таблицы')
            self.status.emit('Создаем шаблон таблицы')
            self.progress.emit(progress)
            document = docx.Document()  # Открываем
            style = document.styles['Normal']
            font = style.font
            font.name = 'Time New Roman'
            font.size = Pt(11)
            section = document.sections[0]
            section.top_margin = Cm(1)  # Верхний отступ
            section.bottom_margin = Cm(1)  # Нижний отступ
            section.left_margin = Cm(1)  # Отступ слева
            section.right_margin = Cm(1)  # Отступ справа
            section.orientation = WD_ORIENTATION.LANDSCAPE  # Альбомная ориентация
            section.page_width = Cm(42)
            section.page_height = Cm(29.7)
            table = document.add_table(rows=5, cols=18, style='Table Grid')
            table.autofit = True
            table.cell(2, 0).merge(table.cell(2, 17))
            table.cell(4, 0).merge(table.cell(4, 17))
            for i in range(18):
                if i < 9 or i in [12, 13]:
                    table.cell(0, i).merge(table.cell(1, i))

                    table.cell(0, i).text = name_col[i]
                elif i == 9:
                    table.cell(0, i).merge(table.cell(0, 11))
                    table.cell(1, i).text = name_col[i]
                    table.cell(0, i).text = name_1_col[0]
                elif i in [10, 11]:
                    table.cell(1, i).text = name_col[i]
                elif i == 14:
                    table.cell(0, i).merge(table.cell(0, 15))
                    table.cell(1, i).text = name_col[i]
                    table.cell(0, i).text = name_1_col[1]
                elif i == 16:
                    table.cell(0, i).merge(table.cell(0, 17))
                    table.cell(1, 16).width = Cm(51)
                    table.cell(1, 17).width = Cm(51)
                    table.cell(1, i).text = name_col[i]
                    table.cell(0, i).text = name_1_col[2]
                else:
                    table.cell(1, i).text = name_col[i]
                if i in [0, 1, 6, 7, 8, 12, 13]:
                    self.set_vertical_cell_direction(table.cell(0, i), "btLr")
                elif 8 < i < 12 or i > 13:
                    self.set_vertical_cell_direction(table.cell(1, i), "btLr")
                table.cell(1, i).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                table.cell(0, i).vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                table.cell(1, i).vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            table.cell(0, 9).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            table.cell(0, 14).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            table.cell(0, 16).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            table.cell(2, 0).text = '{%tr for item in table_contents %}'

            # table.rows[3].height_rule = WD_ROW_HEIGHT_RULE.AUTO
            table.rows[1].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            table.rows[1].height = Cm(8)
            table.rows[3].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            table.rows[3].height = Cm(5)
            for i, j in enumerate(index):
                table.cell(3, i).text = '{{item.' + j + '}}'
                table.cell(3, i).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                table.cell(3, i).vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                if i > 11:
                    self.set_vertical_cell_direction(table.cell(3, i), "btLr")

            table.cell(4, 0).text = '{%tr endfor %}'
            context = {'table_contents': table_contents}
            document.save(str(pathlib.Path(self.finish_path, str(self.file_name) + '.docx')))
            self.logging.info('Заносим данные в таблицу')
            self.status.emit('Заносим данные в таблицу')
            template = DocxTemplate(str(pathlib.Path(self.finish_path, str(self.file_name) + '.docx')))
            template.render(context)
            template.save(str(pathlib.Path(self.finish_path, str(self.file_name) + '.docx')))
            progress += 5
            self.progress.emit(progress)
            self.logging.info('Объединяем номера комплектов, проверяем высоту столбцов и изменяем, если нужно')
            self.status.emit('Изменяем форматирование')
            number_start = False
            number_finish = False
            document = docx.Document(str(pathlib.Path(self.finish_path, str(self.file_name) + '.docx')))
            table = document.tables[0]
            percent = 50 / len(number_set)
            for number in number_set:
                if number_start is False:
                    number_start = number
                elif number_finish is False:
                    number_finish = number - 1
                    self.status.emit(f'Изменяем форматирование строк с {number_start} по {number_finish}')
                    value = table.cell(number_start, 0).text
                    table.cell(number_start, 0).merge(table.cell(number_finish, 0))
                    table.cell(number_start, 0).text = value
                    table.cell(number_start, 0).vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    table.cell(number_start, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    self.set_vertical_cell_direction(table.cell(number_start, 0), "btLr")
                    value_second = table.cell(number_start, 1).text
                    number_start_sec = number_start
                    for numb in range(number_start + 1, number_finish + 1):
                        if table.cell(numb, 1).text == value_second:
                            table.cell(number_start_sec, 1).merge(table.cell(numb, 1))
                            table.cell(number_start_sec, 1).text = value_second
                            table.cell(number_start_sec, 1).vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            table.cell(number_start_sec, 1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        value_second = table.cell(numb, 1).text
                        number_start_sec = numb
                    number_start = number
                    number_finish = False
                if len(number_set) > 100 and number_set.index(number) % 50 == 0:
                    document.save(str(pathlib.Path(self.finish_path, str(self.file_name) + '.docx')))
                    document = docx.Document(str(pathlib.Path(self.finish_path, str(self.file_name) + '.docx')))
                    table = document.tables[0]
                progress += percent
                self.progress.emit(int(progress))
            # progress += 50
            # self.progress.emit(progress)
            if len_rows:
                percent = 10/len(len_rows)
                # document = docx.Document(str(pathlib.Path(self.finish_path, str(self.file_name) + '.docx')))
                # table = document.tables[0]
                for key in len_rows:
                    table.rows[key].height = Cm(len_rows[key])
                    progress += percent
                    self.progress.emit(progress)
            document.save(str(pathlib.Path(self.finish_path, str(self.file_name) + '.docx')))
            os.startfile(pathlib.Path(self.finish_path, f'{self.file_name}.docx'))
            os.startfile(pathlib.Path(self.finish_path, f'{self.file_name}.xlsx'))
            self.logging.info("Конец программы, время работы: " + str(datetime.datetime.now() - time_start))
            self.logging.info("\n*******************************************************************************\n")
            if incoming_errors:
                self.logging.info("Выводим ошибки")
                self.status.emit('Готово, есть ошибки.')  # Посылаем значние если готово
                self.queue.put({'errors': incoming_errors, 'title': 'Ошибки в загруженных данных, номера строк:\n'})
                self.errors.emit()
            else:
                self.status.emit('Готово!')  # Посылаем значние если готово
                self.progress.emit(100)  # Завершаем прогресс бар
        except BaseException as e:  # Если ошибка
            self.status.emit('Ошибка')  # Сообщение в статус бар
            self.logging.error("Ошибка:\n " + str(e) + '\n' + traceback.format_exc())
            self.progress.emit(0)
